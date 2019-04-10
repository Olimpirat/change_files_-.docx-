#!/usr/bin/python3
# -*- coding: utf-8 -*-

import docx, re, os

"""Программа находит в файле .docx (ms word 2007+) текст(искомый фрагмент) и замещает его с сохранением стиля текста.
    Предусмотрена возможность создавать копии файлов и нормальная работа если текст "программно" разбит на куски MS WORD. 
   Тест работоспособности проводился на двух машинах с Windows 10. С GUI на Windows XP не работает(скорей всего виновата сборка в исполняющий файл)"""

#ПО ХОРОШЕМУ, СДЕЛАТЬ ВОЗМОЖНОСТЬ РАБОТЫ НЕ ТОЛЬКО НА WINDOWS. В первую очередь работа с путями (в частности с копированием файлов)

def refind_in_paragraphs(doc_parag, text_find, text_refind, list_repl=list()): #поиск и замена в параграфе (или ячейки таблицы)
    list_of_changes = list_repl
    for x in doc_parag.paragraphs: #обходим параграфы текста
        if text_find in x.text: #если есть, то выполняем замену.
            substitution_values = list()  # значения под замену
            replacement_positions = list(m.start() for m in re.finditer(text_find, x.text)) #исходные позиции замен

            for para in list_of_changes: #если список найденных попадает в список уже заменненых, то пропустить все замены
                if x.text == para[0]:
                    if replacement_positions[0] in range(para[1][0][0], para[1][0][1]): #проверка была ли замена этого значения, если ДА, то пропистуть
                        return None
                #иначе уходит на замену

            number_run = 0  # номер прогона
            inline = x.runs # полученный текст с прогона
            for i in range(len(inline)): #разбираем построчно
                #print("строка №" + str(i) + "\n")
                #print(inline[i].text)

                if text_find in inline[i].text: #проверка каждого слова и замена всех совпадений
                    text = inline[i].text.replace(text_find, text_refind)
                    inline[i].text = text
                    #print("сделано в параграфах! " + text)
                else: #текст на прогонах разделен внутренними разделителями ворда, связанными с исправлениями текста. Далее идет работа по кусочкам нужного текста

                    if number_run != 0:
                        start = stop
                        stop = len(inline[i].text) + start
                        if ref_stop == None:
                            break  # если совершен последний прогон под замену, тогда остановиться
                        ref_start = ref_stop
                    else:
                        start = x.text.find(text_find)
                        stop = len(inline[i].text)
                        ref_start = 0
                        number_run += 1
                    ref_stop = stop - start + ref_start
                    try:
                        if text_find[-(len(text_find) - ref_start):] == inline[i].text.split()[0]:  # если концовка искомого текста, тогда:
                            ref_stop = None
                    except IndexError: #если "перелет", то ничего не делать и просто продолжить без ошибки
                        pass

                    text = inline[i].text.replace(text_find[ref_start:ref_stop], text_refind[ref_start:ref_stop]) #создаем кусок под замену
                    inline[i].text = text #заменяем кусок
                    #print("сделано в параграфах! " + text)

            substitution_values.append(x.text) #добавляем строку параграфа
            substitution_values.append(change_numbers(replacement_positions,text_find, text_refind)) #добавляем кортеж с позициями замен
            list_of_changes.append(substitution_values) #добавляем кортеж с текстом параграфа и вложеным кортежем с номерами замен

def change_numbers(list_mail, text_find, text_refind): #функция принимает список исходных замен и возвращает кортеж из кортежей замен от-до
    list_for_return = list()
    len_difference = len(text_refind) - len(text_find) #разница кол-ва символов между заменяемым и новым
    number = 0
    for x in list_mail:
        from_number = x + len_difference * number
        before_number = from_number + len(text_refind)
        replace_tup = (from_number, before_number)
        list_for_return.append(replace_tup)
        list_for_return.append(text_find) #добавляем значения замен, чтобы одно на одно менялось
        list_for_return.append(text_refind)
        number += 1
    return tuple(list_for_return)

def refind_in_table(doc_table, text_find, text_refind, list_repl): #поиск по таблице, с "ячейкой" работаем как с параграфом
    for y in doc_table.tables:  # обходим таблицу
        for ro in y.rows:  # обходим строки таблицы
            for roc in ro.cells:  # обходим ячейки в строку
                if text_find in roc.text:
                    refind_in_paragraphs(roc, text_find, text_refind, list_repl)

def find_and_replace_text(path, dict_f_rf, copy=False): #Поиск и замена текста
    try:
        document = docx.Document(r'{0}'.format(path)) #открываем документ
    except docx.opc.exceptions.PackageNotFoundError: #если файл не корректный или поврежден
       #print("The file is incorrect or damaged.")
       raise FileNotFoundError  #вывод сообщения пользователю и работа в "интерфейсе пользователя"

    list_repl=list() # создаем список, куда будем помещать кортежи с уже измененными данными

    if copy: #Если надо сделать копию, то создаем в директории папку copy и туда скидываем документы. Работать должно только на Windows
        point = path.rfind('\\')
        if not os.path.isdir(r"{0}{1}".format(path[:point], "\\copy")):
            os.mkdir(r"{0}{1}".format(path[:point], "\\copy"))
        path_copy = r"{0}{1}{2}".format(path[:point], "\\copy", path[point:])
        document.save(r'{0}'.format(path_copy))

    for text_find, text_refind in dict_f_rf.items(): #вытаскиваем что и на что надо менять и начинаем искать
        refind_in_paragraphs(document, text_find, text_refind, list_repl) #ищем по параграфам
        refind_in_table(document, text_find, text_refind, list_repl) #ищем по таблицам

    try:
        document.save(r'{0}'.format(path))
    except PermissionError: #Если изменненый файл нельзя закрыть (скорей всего файл открыт)
        raise PermissionError #вывод пользователю сообщения и последующая работа в "интерфейсе пользователя"

        #старый консольный кусок
        #print("Закройте сначало файл куда сохраняются данные!!! Потом нажмите Enter")
        #input("Нажмите Enter для продолжения")
        #find_and_replace_text(path, dict_f_rf, copy=False)

if __name__ == "__main__":
    #path = r"H:\работа\копия\Некрасовка съем\Справки октябрь 2018\ОДС-3\demo.docx"
    #dict_test = {'Февраль': 'Январь','Январь': 'xasxasxas', '01.11.2018': '01.12.2018', '01.12.2018': '01.01.2019'}
    path = r"H:\работа\копия\Некрасовка съем\Справки октябрь 2018\ОДС-1\типо докс.docx"
    dict_test = {'Февраль': 'Январь'}
    find_and_replace_text(path, dict_test, copy=True)
